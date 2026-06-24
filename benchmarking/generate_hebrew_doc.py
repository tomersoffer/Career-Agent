# -*- coding: utf-8 -*-
"""
Generates the Hebrew (.docx) decision-log documentation for the
Smart Career Navigator data-engineering project (Phase 1).
RTL-formatted throughout.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "David"
HEAD_FONT = "Arial"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # dark blue
GREY = RGBColor(0x55, 0x55, 0x55)


# ---------- RTL helpers ----------
def _set_rtl_paragraph(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _style_run(run, font=BODY_FONT, size=11, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    # complex-script font (Hebrew)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), font)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    # mark run as RTL
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)


def add_par(doc, text, size=11, bold=False, color=None, font=BODY_FONT, space_after=6, space_before=0):
    p = doc.add_paragraph()
    _set_rtl_paragraph(p)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    _style_run(run, font=font, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 22, 1: 16, 2: 13}
    colors = {0: ACCENT, 1: ACCENT, 2: RGBColor(0x2E, 0x6B, 0x2E)}
    p = doc.add_paragraph()
    _set_rtl_paragraph(p)
    p.paragraph_format.space_before = Pt(14 if level else 0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _style_run(run, font=HEAD_FONT, size=sizes.get(level, 13), bold=True, color=colors.get(level, ACCENT))
    return p


def add_bullet(doc, text, sub=False):
    p = doc.add_paragraph()
    _set_rtl_paragraph(p)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.right_indent = Inches(0.3 if not sub else 0.6)
    bullet = "•  " if not sub else "–  "
    run = p.add_run(bullet + text)
    _style_run(run, size=11)
    return p


def add_kv_bullet(doc, key, val, sub=False):
    """Bullet with a bold key and normal value."""
    p = doc.add_paragraph()
    _set_rtl_paragraph(p)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.right_indent = Inches(0.3 if not sub else 0.6)
    r1 = p.add_run(("•  " if not sub else "–  ") + key + " — ")
    _style_run(r1, size=11, bold=True)
    r2 = p.add_run(val)
    _style_run(r2, size=11)
    return p


def add_rtl_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # RTL visual order for the table
    tblPr = table._tbl.tblPr
    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0]._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hdr[i].paragraphs[0].add_run(h)
        _style_run(run, font=HEAD_FONT, size=10, bold=True, color=RGBColor(255, 255, 255))
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            para = cells[i].paragraphs[0]
            para._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(str(val))
            _style_run(run, size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    return table


def add_spacer(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)


# ---------- flow-diagram helpers ----------
def _par_shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def _par_border(p, color="1F4E79", sz=6):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)


def add_flow_box(doc, text, fill="DCE6F1"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.left_indent = Inches(0.7)
    p.paragraph_format.right_indent = Inches(0.7)
    _par_shade(p, fill)
    _par_border(p)
    run = p.add_run(text)
    _style_run(run, font=HEAD_FONT, size=10.5, bold=True, color=ACCENT)
    return p


def add_flow_arrow(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("▼")
    _style_run(run, font=HEAD_FONT, size=12, bold=True, color=GREY)
    return p


def add_lead(doc, text):
    """A bold accented sub-sub-heading inside a subsection."""
    return add_par(doc, text, size=12, bold=True, color=ACCENT, font=HEAD_FONT,
                   space_before=8, space_after=3)


def add_code_block(doc, code, size=8.5):
    """LTR monospace pseudocode block with light shading; preserves indentation."""
    for raw in code.split("\n"):
        stripped = raw.lstrip(" ")
        indent = len(raw) - len(stripped)
        text = (" " * indent + stripped) if stripped else " "
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.12)
        _par_shade(p, "F3F3F3")
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), "Consolas")
        rFonts.set(qn('w:hAnsi'), "Consolas")
        rFonts.set(qn('w:cs'), "Consolas")
        rPr.append(rFonts)
    return


# ====================================================================
doc = Document()
# default style
normal = doc.styles['Normal']
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(0.9)

# ---------- TITLE ----------
add_heading(doc, "תיעוד תהליך הנדסת הנתונים", level=0)
add_par(doc, "פרויקט גמר — Smart Career Navigator", size=15, bold=True, color=GREY, font=HEAD_FONT, space_after=2)
add_par(doc, "סוכן AI להמלצת משרות על בסיס תפקיד, ניסיון ומיקום גאוגרפי", size=12, color=GREY, space_after=2)
add_par(doc, "שלב 1: עיבוד מקדים והכנת מאגר הנתונים המאוחד (Master Table)", size=11, color=GREY, space_after=14)

# ---------- 1. INTRO ----------
add_heading(doc, "1. מבוא ורקע הפרויקט", level=1)

add_heading(doc, "1.1 רקע על הדומיין (כללי)", level=2)
add_par(doc, "תחום הפרויקט הוא שוק העבודה והגיוס הדיגיטלי (HR-Tech). פלטפורמות תעסוקה מקוונות כדוגמת "
             "LinkedIn מייצרות מדי יום היקף עצום של מודעות דרושים — טקסט חופשי ולא-מובנה (כותרת, תיאור, "
             "דרישות וכישורים) לצד שדות מובנים כגון מיקום, רמת ניסיון ושכר. האתגר המרכזי בדומיין הוא התאמה "
             "(Matching) בקנה מידה גדול בין מועמד לבין המשרה הנכונה עבורו, מתוך מאות אלפי אפשרויות. בשנים "
             "האחרונות הופכים כלי למידת מכונה (ML) ועיבוד שפה טבעית (NLP) למנוע המרכזי של התחום, ומאפשרים "
             "מעבר מחיפוש מבוסס מילות-מפתח להבנה סמנטית של כוונת המשתמש ושל תוכן המשרה.")

add_heading(doc, "1.2 מהות הפרויקט ובחירת הנושא (ספציפי)", level=2)
add_par(doc, "בחרנו בנושא משתי סיבות. ראשית, רלוונטיות אישית: כסטודנטים בפתח הקריירה, מציאת המשרה המתאימה "
             "היא בעיה אמיתית ומוחשית עבורנו. שנית, הזדמנות לימודית: מאגר \"LinkedIn Job Postings\" מציע "
             "נתוני אמת עשירים בהיקף Big Data (123,849 משרות), המשלבים טקסט חופשי ושדות מובנים — קרקע "
             "אידיאלית ליישום מגוון הטכניקות שנלמדו בקורס. מהות הפרויקט היא בניית סוכן בינה מלאכותית, "
             "\"Smart Career Navigator\", המבצע פעולת התאמה פרואקטיבית (ולא רק מציג רשימת תוצאות) על בסיס "
             "שלושה צירים: תפקיד (Role), רמת ניסיון (Experience) ומיקום גאוגרפי (Geography). הסוכן משלב "
             "מודלים קלאסיים של דמיון (Cosine / Jaccard) ואשכול (K-Means) עם ממשק שיחה בשפה טבעית מבוסס LLM.")

add_heading(doc, "1.3 הצהרת הבעיה והיעדים", level=2)
add_par(doc, "מחפש העבודה ניצב בפני עומס מידע: מאות אלפי מודעות, חיפוש מבוסס מילות-מפתח המחזיר תוצאות רועשות, "
             "נתוני שכר חסרים ומשובשים, וקושי להעריך בו-זמנית התאמה לפי תפקיד, ניסיון ומיקום. בנוסף, חסרה "
             "הכוונה יזומה המתריעה על מודעות חשודות או מסייעת בשיפור המועמדות.")
add_par(doc, "היעדים שהסוכן שואף להשיג:", bold=True, space_before=2)
add_bullet(doc, "להבין בקשה חופשית בשפה טבעית של המשתמש.")
add_bullet(doc, "לאחזר את המשרות הרלוונטיות ביותר לפי תפקיד / ניסיון / מיקום.")
add_bullet(doc, "להתריע על משרות חריגות (אנומליות) כפרופיל סיכון.")
add_bullet(doc, "לענות על שאלות מבוססות-נתונים אמיתיים מעל המאגר.")
add_bullet(doc, "לספק ייעוץ קריירה והתאמת קורות-חיים (CV) לתפקיד.")
add_bullet(doc, "לשמור על שקיפות ויושרה — כל מספר מחושב דטרמיניסטית וה-ML מדיד וניתן להסבר.")

add_heading(doc, "1.4 הגדרת הבעיה ורלוונטיותה לאתגרי NLP/LLM הנוכחיים", level=2)
add_par(doc, "מנקודת מבט מחקרית, הבעיה היא התאמה סמנטית בין שאילתה בשפה טבעית לבין רשומות מובנות — לב האתגר "
             "של מערכות אחזור מבוסס-יצירה (RAG). הפרויקט נוגע במספר אתגרי ליבה עכשוויים של NLP/LLM:")
add_kv_bullet(doc, "עיגון ומניעת הזיות (Grounding & Hallucination)", "מודלי שפה נוטים \"להמציא\" עובדות "
              "ומספרים; אנו מתמודדים עם זה באמצעות מנוע אנליטיקה ב\"רשימת היתר\" המחשב כל מספר ב-Pandas, "
              "ובעיגון דרישות המשרה (must-haves) באופן דטרמיניסטי — ה-LLM מנסח בלבד.")
add_kv_bullet(doc, "ארכיטקטורה היברידית", "שילוב ML קלאסי (TF-IDF, דמיון קוסינוס, K-Means) עם LLM מפצה על "
              "חולשת המודל בחישוב מספרי ובדיוק עובדתי, ומותיר לו את שפת ההבנה והניסוח שבהם הוא מצטיין.")
add_kv_bullet(doc, "רב-לשוניות", "הפקת תשובות בעברית מתוך מאגר נתונים באנגלית.")
add_kv_bullet(doc, "שקיפות ויושרה אקדמית", "העדפת רכיבים מדידים וניתנים-להסבר על פני \"קופסה שחורה\", "
              "תוך שמירה על עלות וזמן-תגובה סבירים (RAG בזיכרון, ללא מסד נתונים וקטורי חיצוני).")

add_heading(doc, "1.5 חדשנות ומקוריות", level=2)
add_par(doc, "החדשנות של הפתרון אינה בשימוש ב-LLM כשלעצמו, אלא באופן שבו הוא משולב עם מודלים קלאסיים של "
             "למידת מכונה כדי לתת לבעיה מענה מדיד, כן ועמיד-להזיות. עיקרי החדשנות בפתרון המוצע:")
add_kv_bullet(doc, "ארכיטקטורה היברידית אמיתית", "ה-ML (TF-IDF, קוסינוס, Jaccard, K-Means) מבצע את העבודה "
              "המדידה — אחזור, דירוג ואשכול — וה-LLM מנסח ומשוחח בלבד. כך כל החלטה ניתנת להסבר ולהערכה, "
              "להבדיל מ\"עטיפת LLM\" כקופסה שחורה.")
add_kv_bullet(doc, "התאמה פרואקטיבית עם פרופיל סיכון", "מעבר להצגת משרות, הסוכן מזהה אנומליות שוק (מרחק "
              "אוקלידי מצנטרואיד האשכול) ומתריע בשפה פיננסית על מודעות חריגות/חשודות — מסגור ההתאמה גם "
              "כשירות הערכת-סיכון.")
add_kv_bullet(doc, "אנליטיקה עמידה-להזיות", "מנוע \"רשימת היתר\" המחשב כל מספר ב-Pandas (map→compute→narrate), "
              "כך שה-LLM לעולם אינו ממציא נתון; ועיגון דרישות (must-haves) דטרמיניסטי בהתאמת קורות-החיים.")
add_kv_bullet(doc, "התאמת קו\"ח מבוססת-תשאול (Elicitation-first)", "הסוכן קודם מתשאל ניסיון אמיתי שלא תועד, "
              "ורק אז כותב מחדש פעם אחת — ללא בדיית עובדות ועם ניהול פערים כן.")
add_kv_bullet(doc, "ממשק עברית מלא ו-RAG בזיכרון", "שיחה בשפה טבעית בעברית מעל נתונים באנגלית, ללא מסד נתונים "
              "וקטורי חיצוני, עם נתב כוונות יחיד המנתב חיפוש / שאלת-נתונים / ייעוץ ללא קריאת-מודל נוספת.")

add_heading(doc, "1.6 עומק טכני — מושגי NLP/LLM שיישמנו", level=2)
add_par(doc, "הפרויקט נשען על מושגים וטכנולוגיות מליבת ה-NLP וה-LLM שנלמדו בכיתה, ומיישם כל אחד מהם ישירות "
             "על נתוני המשרות שלנו. הטבלה מתמצתת את ההקבלה בין המושג כפי שנלמד לבין היישום הקונקרטי על "
             "המאגר; הפירוט המתמטי המלא מופיע בפרקים 10–12.")
add_rtl_table(doc,
    ["המושג / הטכנולוגיה (כפי שנלמד)", "היישום על הנתונים שלנו"],
    [
        ["ייצוג טקסט וקטורי (TF-IDF)",
         "המרת השדה text_blob (כותרת + כישורים + תיאור) לווקטורים משוקללים — מילה מקבלת משקל לפי תדירותה "
         "במשרה ביחס לנדירותה בכלל המאגר; הבסיס לכל ההתאמה הסמנטית."],
        ["דמיון קוסינוס (Cosine Similarity)",
         "דירוג המשרות לפי קוסינוס הזווית בין וקטור בקשת המשתמש לבין וקטור הכותרת — לכידת התאמה סמנטית "
         "גם כשאין חפיפת מילים מדויקת."],
        ["דמיון ז'קארד (Jaccard Similarity)",
         "חיתוך קבוצת הכישורים של המשתמש מול קבוצת כישורי המשרה, |A∩B|/|A∪B| — מדד מתאים לתגיות בדידות."],
        ["עיבוד מקדים של טקסט (Text Preprocessing)",
         "טוקניזציה, המרה לאותיות קטנות, הסרת מילות עצירה וניקוי URL/HTML/תווי קידוד פגומים — בניית "
         "text_blob נקי ומוכן ל-TF-IDF."],
        ["אשכול לא-מפוקח (K-Means) + תקנון + שיטת המרפק",
         "קיבוץ המשרות ל-5 \"פרופילי שוק\" על מאפיינים מספריים מתוקננים (StandardScaler); בחירת K "
         "אובייקטיבית בשיטת המרפק (Elbow)."],
        ["זיהוי אנומליות (מרחק מצנטרואיד)",
         "סימון משרות חריגות לפי מרחק אוקלידי גדול מאחוזון 95 ממרכז האשכול — בסיס לפרופיל הסיכון של הסוכן."],
        ["RAG — אחזור מבוסס-יצירה",
         "אחזור מקומי בזיכרון של שלוש המשרות הרלוונטיות והזרקתן כהקשר ל-LLM — תשובה מעוגנת בנתונים "
         "אמיתיים, ללא מסד נתונים וקטורי חיצוני."],
        ["LLM והנדסת פרומפטים (Prompt Engineering)",
         "הנחיות מערכת, פלט JSON מובנה, נתב כוונות, ועיגון דטרמיניסטי (whitelist) למניעת הזיות; ניסוח "
         "התשובות בעברית באמצעות מודל Claude."],
    ],
    col_widths=[2.5, 4.2])

add_heading(doc, "1.7 הנתונים והיקף ה-Big Data", level=2)
add_par(doc, "מקור הנתונים: מאגר \"LinkedIn Job Postings (2023–2024)\" מ-Kaggle, המורכב מ-11 קובצי CSV "
             "הקשורים זה לזה (Relational Big Data) סביב שני מפתחות — job_id ו-company_id. נקודת הפתיחה: "
             "123,849 משרות ו-24,473 חברות; לאחר העיבוד התקבלה טבלת אב אחת בת 119,117 משרות — מעל סף "
             "ה-Big Data של 100,000 רשומות.")
add_par(doc, "שיקולים לאיסוף הנתונים:", bold=True, space_before=2)
add_bullet(doc, "היקף: מעל 100K רשומות אמיתיות — עומד בדרישת ה-Big Data של הפרויקט.")
add_bullet(doc, "עושר: שילוב טקסט חופשי (תיאורי משרה) עם שדות מובנים (שכר, ניסיון, מיקום) — מתאים גם "
                "ל-NLP וגם ל-ML מספרי.")
add_bullet(doc, "מבנה רלציוני: 11 טבלאות עם מפתחות וקשרי 1-לרבים — מאפשר תרגול מיזוג והנדסת נתונים אמיתית.")
add_bullet(doc, "זמינות וחוקיות: מאגר ציבורי ב-Kaggle, עדכני (2023–2024) ורלוונטי לשוק העבודה.")
add_par(doc, "אופן האיסוף והאחסון:", bold=True, space_before=2)
add_bullet(doc, "האיסוף: הורדת 11 הקבצים מ-Kaggle ואחסונם מקומית בתיקיית הקלט הגולמי (Job Posting Dataset/).")
add_bullet(doc, "האחסון: התוצר נשמר כקובץ Master CSV ב-UTF-8 (לשימור תווים מיוחדים) + גרסה מצומצמת + "
                "data_dictionary.json; בזמן ריצה הנתונים נטענים לזיכרון (In-Memory) ללא מסד נתונים חיצוני.")
add_par(doc, "הכנה ועיבוד מקדים — תרשים זרימה:", bold=True, space_before=4)
add_flow_box(doc, "מקור: 11 קובצי CSV מ-Kaggle — 123,849 משרות גולמיות")
add_flow_arrow(doc)
add_flow_box(doc, "פרופיילינג: ניתוח חוסרים, חריגים, כפילויות ויחסי 1-לרבים")
add_flow_arrow(doc)
add_flow_box(doc, "החלטות ברמת קובץ: שמירה / איחוד / הסרה (11 קבצים ← טבלת אב)")
add_flow_arrow(doc)
add_flow_box(doc, "ניקוי עמודות: נרמול שכר, פיצול מיקום, קידוד ניסיון/גודל, בניית text_blob")
add_flow_arrow(doc)
add_flow_box(doc, "סינון שורות וחריגים: כפילויות, כותרות פסולות, משרות זרות (−4,732)")
add_flow_arrow(doc)
add_flow_box(doc, "מיזוג בטוח: צמצום 1-לרבים לשורה למפתח + left-join עם validate")
add_flow_arrow(doc)
add_flow_box(doc, "תוצר: Master Table — 119,117 × 22 (CSV ב-UTF-8) + data_dictionary.json", fill="D6E8D6")
add_par(doc, "האם הנתונים נוקו והוכנו מראש? כן — לפני כל שלב ניתוח ומידול בוצע עיבוד מקדים מלא. כל שלב "
             "בתרשים מתועד בהרחבה בפרקים הבאים: פרופיילינג (פרק 2), החלטות קובץ (פרק 4), ניקוי עמודות "
             "(פרק 5), סינון חריגים (פרקים 6–7), ובטיחות המיזוג (פרק 9).", space_before=6)

add_heading(doc, "1.8 הסבר על המימוש: בחירת מודלים, השוואה ופסאודו-קוד", level=2)
add_par(doc, "פרק זה מסביר את המימוש עצמו: אילו מודלים נבחרו וההיגיון מאחורי הבחירה, התאמתם למשימה והיעילות "
             "הפוטנציאלית שלהם, השוואה בין אופציות קיימות, דוגמאות תוצאות על מקבץ קטן של נתונים, ולבסוף "
             "פסאודו-קוד רחב של הפתרון מקצה לקצה (הסבר רעיוני — לא כל פרט בו ממומש כלשונו).")

add_lead(doc, "א. בחירת המודלים וההיגיון מאחוריהם")
add_kv_bullet(doc, "ייצוג ואחזור טקסט — TF-IDF + Cosine Similarity",
              "המרת כותרות המשרה וה-text_blob לווקטורים ודירוג לפי דמיון קוסינוס. נבחר משום שהוא דטרמיניסטי, "
              "שקוף וניתן-להסבר, אינו דורש נתוני אימון או GPU, ויעיל במיוחד על כותרות קצרות — בדיוק מה שנדרש "
              "מ\"ML מדיד\" בקורס.")
add_kv_bullet(doc, "אשכול לא-מפוקח — K-Means",
              "קיבוץ המשרות לפרופילי שוק על מאפיינים מספריים מתוקננים. נבחר בזכות מדרגיות ל-119K שורות, "
              "צנטרואידים ניתנים-לפרשנות, וזיהוי אנומליות טבעי לפי מרחק מהמרכז.")
add_kv_bullet(doc, "שכבת שפה — LLM (Claude Sonnet 4.6)",
              "מודל השפה משמש אך ורק לפענוח הבקשה ולניסוח התשובה בעברית — לא לחישוב. נבחר Claude בזכות "
              "עברית איכותית, פלט JSON מובנה ואמין, ציות-להנחיות חזק (קריטי לעיגון ולמניעת הזיות), "
              "ותמיכה ב-prompt-caching להוזלת עלות.")

add_lead(doc, "ב. התאמה למשימה ויעילות פוטנציאלית")
add_bullet(doc, "TF-IDF + Cosine: המטריצה נבנית פעם אחת ונשמרת; בזמן ריצה האחזור הוא הכפלת וקטורים בזיכרון — "
                "תת-שנייה על 119K משרות, ללא עלות הסקה.")
add_bullet(doc, "K-Means: מאומן פעם אחת לא-מקוון ונשמר (joblib); בזמן ריצה רק שיוך לצנטרואיד הקרוב — זניח.")
add_bullet(doc, "LLM: 1–2 קריאות API לכל תור (פענוח + ניסוח). העלות זניחה; ההשהיה היא הרכיב היקר — ולכן "
                "האחזור מתבצע מקומית (RAG) וה-LLM רואה רק 3 משרות, לא את כל המאגר.")

add_lead(doc, "ג. השוואת אופציות — מדוע נבחר מה שנבחר")
add_par(doc, "השוואה 1 — ייצוג טקסט לאחזור:", bold=True, space_before=2)
add_rtl_table(doc,
    ["קריטריון", "TF-IDF + Cosine (נבחר)", "Dense Embeddings (Sentence-BERT)"],
    [
        ["עיקרון", "שכיחות מונחים משוקללת + קוסינוס", "וקטורים צפופים ממודל מאומן מראש"],
        ["אימון / חומרה", "אין — נבנה ישירות מהמאגר", "מודל כבד; הסקה מצריכה GPU/שירות חיצוני"],
        ["פרשנות והסבר", "גבוהה — ניתן להצביע על המילים המשפיעות", "נמוכה — וקטור סמוי (קופסה שחורה)"],
        ["כותרות קצרות", "מצוין — חפיפת מונחים ישירה", "היתרון פוחת על טקסט קצר מאוד"],
        ["מהירות / עלות", "מהיר מאוד, מקומי, ללא עלות", "איטי ויקר יותר"],
        ["התאמה ל\"ML מדיד\"", "מלאה", "חלקית"],
    ],
    col_widths=[1.6, 2.6, 2.5])
add_par(doc, "מסקנה: נבחר TF-IDF + Cosine — שקיפות, יעילות והתאמה לדרישות הקורס. Embeddings צפופים נשמרים "
             "כשדרוג עתידי אפשרי.", space_before=2)
add_par(doc, "השוואה 2 — אלגוריתם אשכול:", bold=True, space_before=4)
add_rtl_table(doc,
    ["קריטריון", "K-Means (נבחר)", "Hierarchical / DBSCAN"],
    [
        ["מדרגיות ל-119K", "מצוינת (כמעט לינארי)", "חלשה (היררכי ~O(n²)) / DBSCAN רגיש לפרמטרים"],
        ["פרשנות", "צנטרואידים = פרופילי שוק ברורים", "דנדרוגרם כבד / אשכולות בצורה חופשית קשים לפרשנות"],
        ["זיהוי אנומליות", "טבעי — מרחק מצנטרואיד", "פחות ישיר"],
        ["בחירת מספר קבוצות", "שיטת המרפק (אובייקטיבית)", "חיתוך דנדרוגרם / eps לא יציב"],
    ],
    col_widths=[1.6, 2.4, 2.7])
add_par(doc, "מסקנה: נבחר K-Means — מדרגי, מפורש ומשתלב ישירות עם שכבת זיהוי האנומליות.", space_before=2)
add_par(doc, "באשר ל-LLM: נבחר Claude Sonnet 4.6 על פני חלופות (כגון GPT / Gemini ועל פני שכבות Opus/Haiku) "
             "כנקודת איזון בין איכות עברית, אמינות פלט JSON, השהיה ועלות — תוך שמירה על תפקיד מצומצם (פענוח "
             "וניסוח בלבד) כדי שהדיוק העובדתי יישאר באחריות ה-ML הדטרמיניסטי.", space_before=2)

add_lead(doc, "ד. דוגמאות תוצאות על מקבץ קטן של נתונים")
add_par(doc, "דוגמה ייצוגית — שאילתה: \"Senior Data Analyst position in NY\". הפענוח מזהה role=data analyst, "
             "state=NY, seniority=mid-senior; מוחל סינון גאוגרפי קשיח ל-NY, ושלוש המשרות המדורגות גבוה ביותר:")
add_rtl_table(doc,
    ["דירוג", "כותרת המשרה", "מדינה", "ציון התאמה", "סרגל שכר", "אנומליה"],
    [
        ["1", "Senior Data Analyst", "NY", "0.91", "5 ($100K–150K)", "לא"],
        ["2", "Data Analyst, Analytics", "NY", "0.78", "4 ($80K–100K)", "לא"],
        ["3", "Lead Business Data Analyst", "NY", "0.66", "6 ($150K–200K)", "כן ⚠️"],
    ],
    col_widths=[0.7, 2.6, 0.8, 1.0, 1.4, 0.9])
add_bullet(doc, "התרעת סיכון: המשרה השלישית סומנה is_anomaly=1 (שכר חריג ביחס לפרופיל האשכול), והסוכן צירף "
                "אזהרה יזומה למשתמש לאמת את הנתונים.")
add_bullet(doc, "דוגמת שאלת-נתונים: \"מה השכר הטיפוסי ל-Data Analyst?\" — מנוע האנליטיקה מחשב חציון (למשל "
                "~$85K, n משרות) ב-Pandas, וה-LLM מנסח את המספר בעברית בלבד.")
add_bullet(doc, "דוגמת אשכול: על מדגם של 5 משרות, כל אחת שויכה לאחד מ-5 פרופילי השוק (פרק 10) עם מרחק "
                "מהצנטרואיד; המשרות מעל אחוזון 95 סומנו כאנומליות.")

add_lead(doc, "ה. פסאודו-קוד מקצה לקצה")
add_par(doc, "להלן תיאור רעיוני רחב של הפתרון — הלוגיקות, התנאים והבעיות שטופלו — משלב בניית המאגר ואימון "
             "המודלים (לא-מקוון) ועד טיפול בתור שיחה בודד (מקוון):")
add_code_block(doc, """# ========== OFFLINE: DATA ENGINEERING + MODEL TRAINING ==========
function build_master(raw_csvs):
    postings = load("postings.csv")            # 123,849 rows
    # reduce every one-to-many table to ONE row per key  (prevents row explosion)
    skills   = aggregate(job_skills + skills_dict)   -> comma list per job_id
    industry = first_valid(job_industries + industries_dict) per job_id
    company  = clean(companies); size = impute_from(employee_counts)
    master   = left_join(postings, [skills, industry, company], validate="m:1")
    assert rows(master) == rows(postings)      # PROBLEM guarded: merge blow-up

    # ---- salary: smart multi-cause recovery (FIX, don't drop) ----
    for job in master:
        raw    = med_salary  or  mean(min, max)
        annual = normalize(raw, pay_period)    # hourly*2080, weekly*52, monthly*12 ...
        annual = recover_mislabeled(annual, raw)   # YEARLY<50 -> hourly; 50..999 -> *1000
        if fulltime and annual < MIN_WAGE: drop_row(job)      # CONDITION: junk salary
        job.salary_band = bucketize(annual)    # 0..7 ; missing -> band 0

    master.text_blob = clean_text(title + skills + description)  # lower, strip url/html, stopwords
    master = filter_usa_only(by job_location)
    master = drop_duplicates(keep = most_viewed)        # 4,673 reposts collapse
    save(master, encoding="utf-8"); save("data_dictionary.json")
    return master                              # 119,117 rows x 22 cols

function train_models(master):
    X       = StandardScaler([experience_rank, company_size, salary_band, remote_allowed])
    K       = elbow(X, range=2..10)            # objective: max distance from chord -> K=5
    kmeans  = KMeans(K).fit(X);  master.cluster = kmeans.labels
    master.dist       = euclidean(X, centroid[label])
    threshold         = percentile(master.dist, 95)
    master.is_anomaly = master.dist > threshold          # ~5% flagged
    tfidf        = TfidfVectorizer().fit(master.title)
    title_matrix = tfidf.transform(master.title)
    save(kmeans, scaler, tfidf, title_matrix)

# ========== ONLINE: ONE AGENT TURN ==========
function handle_turn(prompt, history, cv_profile, screen_matches):
    if no_api_key: return raw_search(prompt)             # graceful degrade, no LLM

    intent = parse_query(prompt)     # 1 LLM call -> {role, synonyms[], city, state,
                                     #   nearby_states[], years, seniority, mode}
    intent.state = clamp_to_valid_states(intent.state)   # PROBLEM: LLM drift -> ground

    switch intent.mode:              # intent router (NO extra LLM call)
        case "data_question": return answer_data_question(prompt)
        case "advice":        return advice_reply(prompt, history, screen_matches, cv_profile)
        case "smalltalk":     return small_talk(prompt)
        case "more":          return next_page(last_results)
        default ("search"):   return search_and_reply(intent, cv_profile)   # also if ambiguous

function search_and_reply(intent, cv_profile):
    pool = master
    if intent.state: pool = pool[ job_state == intent.state ]        # hard geo filter
    # title score = MAX over {role + synonyms}, each scored as its own phrase
    title_s = max_over_variants(cosine(tfidf(variant), title_matrix[pool]))
    exp_s   = experience_proximity(intent.seniority, pool.experience_rank)
    score   = 0.7*title_s + 0.3*exp_s
    cand    = pool[ title_s >= TITLE_FLOOR ]                         # relevance gate

    if empty(cand) and intent.nearby_states:                        # geo fallback cascade
        for st in intent.nearby_states:
            cand = search(st, floor = MIN_FALLBACK)                  # higher bar to justify
            if not empty(cand): break
    if empty(cand): return explain_no_match(intent)                 # honest no-result

    top   = top_n(cand by score, 3)
    facts = [title, company, state, salary_band, experience, is_anomaly, urls]
    reply = synthesize(facts, intent)          # 1 LLM call -> short Hebrew advisor reply
    for job in top:
        if job.is_anomaly: reply.prepend("WARNING risk alert: market anomaly - verify")
    return reply + cards(top)                   # cards carry apply links (passive metadata)

function answer_data_question(q):
    spec = LLM_pick(q, WHITELIST of 7 funcs)    # returns {func, params} via dict lookup ONLY
    rows = master                               #   (no eval / getattr / exec)
    if spec.role: rows = rows[ cosine(role, title) >= 0.30 ]         # ML picks the rows
    value, n = compute[spec.func](rows, spec.params)                 # pandas ONLY, never the LLM
    note = "based on few postings" if n < 5 else ""                  # small-n honesty guard
    return narrate(value, n, note)              # LLM phrases the numbers only""")

add_heading(doc, "1.9 בדיקה ואימות", level=2)
add_par(doc, "איכות הפתרון נבדקה בכמה רבדים — מבדיקות יחידה אוטומטיות ועד בחינה ידנית של איכות התשובות — "
             "במטרה לוודא נכונות, שחזוריות ועמידות במקרי קצה.")

add_lead(doc, "א. שיטות הבדיקה והאימות")
add_kv_bullet(doc, "בדיקות יחידה (pytest)", "כ-100 בדיקות backend ועשרות בדיקות נוספות למודולי ה-CV, על "
              "מודולים טהורים עם הזרקת תלויות (DI) הנבדקים בבידוד.")
add_kv_bullet(doc, "בדיקות אינטגרציה", "Flask test-client מוודא שכל הנתיבים פועלים גם ללא מפתח LLM (smoke), "
              "לצד בדיקות אינטגרציה ייעודיות לנתב הכוונות.")
add_kv_bullet(doc, "אימות הנדסת הנתונים", "השימוש ב-merge(validate=\"m:1\") ובדיקת מספר השורות אחרי כל מיזוג "
              "עוצרים בשגיאה כל ניפוח שורות; כל הספירות אומתו ב-Pandas (ולא ב-wc -l, שמנפח את הספירה פי ~15).")
add_kv_bullet(doc, "אימות המודלים", "בחירת K אובייקטיבית (שיטת המרפק), בדיקת איזון האשכולות, וסף אנומליה "
              "באחוזון 95 שאומת אמפירית (5,931 משרות = 5.0%).")
add_kv_bullet(doc, "עמידות-להזיות", "מנוע האנליטיקה הדטרמיניסטי (whitelist) מפיק תוצאות שחזוריות הניתנות "
              "להשוואה מול חישוב Pandas ידני; בנוסף, עיגון דרישות (must-haves) ומשמר n קטן.")
add_kv_bullet(doc, "בדיקות קצה-לקצה חיות", "הרצת שאילתות אמיתיות בדפדפן עם מפתח API ובחינת איכות ונכונות "
              "התשובה בעברית.")
add_kv_bullet(doc, "טיפול במקרי כשל", "סף רלוונטיות שמחזיר \"אין התאמה\" כן במקום תוצאה שגויה, ודרגרדציה "
              "מבוקרת (הצגת ההקשר הגולמי) כשאין מפתח API.")

add_lead(doc, "ב. דוגמאות תוצאות (לפחות 3 מקרים, כולל מקרה כשל)")
add_rtl_table(doc,
    ["מקרה", "קלט לדוגמה", "התנהגות הסוכן והתוצאה", "אימות"],
    [
        ["1. חיפוש (הצלחה)",
         "\"Senior Data Analyst in NY\"",
         "זוהו role + state + seniority; סינון קשיח ל-NY; הוחזרו 3 משרות אנליסט מדורגות + התרעת אנומליה על אחת.",
         "התוצאות במדינה ובתפקיד הנכונים; ציוני הקוסינוס יורדים בהתאם לדירוג."],
        ["2. שאלת-נתונים (הצלחה)",
         "\"מה השכר הטיפוסי ל-Data Analyst?\"",
         "נותב ל-data_question; חושב חציון ב-Pandas; נוסח בעברית בלבד.",
         "המספר זהה לחישוב Pandas ידני (אימות מוצלב); גודל המדגם n מצוין."],
        ["3. נפילה גאוגרפית (קצה)",
         "\"ML Engineer in Wyoming\"",
         "אין התאמה מעל הסף ב-WY → מפל למדינה סמוכה (CO) → הוחזרה משרת ML Engineer אמיתית ב-CO.",
         "המערכת לא המציאה משרה ב-WY; המעבר הגאוגרפי שקוף ומבוסס קרבה."],
        ["4. כשל בזיהוי הבקשה",
         "\"אני רוצה משהו טוב\" (עמום: ללא תפקיד/מיקום/כוונה)",
         "הכוונה אינה חד-משמעית; ברירת המחדל היא search, אך לא נמצא תפקיד מעל סף הרלוונטיות → הסוכן מחזיר "
         "הודעת \"לא זוהתה בקשה ברורה\" ומבקש חידוד, במקום להמציא תוצאה.",
         "כשל מטופל: לא הוחזרו תוצאות שווא; המשתמש מונחה לנסח מחדש."],
    ],
    col_widths=[1.2, 1.7, 2.4, 1.9])
add_par(doc, "ניתוח מקרה הכשל: סיווג כוונה בקריאה אחת עלול לטעות בבקשות עמומות או רב-כווניות (למשל שאלת-נתונים "
             "המנוסחת באופן חופשי ומסווגת כ-search). זוהי מגבלה ידועה, והיא מוקטנת בשלוש דרכים: ברירת מחדל "
             "בטוחה ל-search (שלעולם אינה גרועה מהבסיס), סף רלוונטיות שמונע תוצאות שווא ומחזיר בקשת חידוד כנה, "
             "ושמירת היסטוריית השיחה כך שהמשתמש יכול לתקן את עצמו בתור הבא. שדרוג עתידי אפשרי: שלב אישור/הבהרה "
             "יזום כאשר ביטחון הסיווג נמוך.", space_before=3)

add_heading(doc, "1.10 שיקולים אתיים", level=2)
add_par(doc, "סוכן הממליץ על משרות נוגע בהחלטות משמעותיות בחיי המשתמש, ולכן נדרשת מודעות להשלכות אתיות — "
             "במיוחד להטיות אפשריות בנתונים ובמודלים, ולשימוש אחראי בבינה מלאכותית.")

add_lead(doc, "א. הטיות בנתונים (Data Bias)")
add_kv_bullet(doc, "הטיה גאוגרפית", "המאגר הוא ארה\"ב בלבד (99.95% מהמשרות); משרות זרות הוסרו. כתוצאה, "
              "הסוכן מוטה לשוק האמריקאי — שוק ישראל נתמך בנפרד דרך הסורק החי, אך הכיסוי אינו סימטרי.")
add_kv_bullet(doc, "הטיית מקור ודגימה", "מודעות LinkedIn נוטות לעבר תפקידי צווארון-לבן, היי-טק ודוברי "
              "אנגלית, ומייצגות-חסר עבודות כפיים, המגזר הלא-פורמלי ומשרות שאינן מתפרסמות בפלטפורמה.")
add_kv_bullet(doc, "הטיית נתונים חסרים", "ל-~70% מהמשרות אין שכר; קידדנו זאת ביושר כ-band 0 (\"לא צוין\") "
              "ולא השלמנו ערכים מלאכותיים — כדי לא להטות סטטיסטיקות שכר ודירוגים.")
add_kv_bullet(doc, "הטיה זמנית", "המאגר הוא תמונת מצב של 2023–2024; שוק העבודה משתנה, וההמלצות עשויות "
              "להתיישן.")

add_lead(doc, "ב. הטיות במודלים (Model Bias)")
add_kv_bullet(doc, "הטיית ייצוג (TF-IDF)", "אחזור מבוסס חפיפת מונחים מעדיף ניסוח שכיח; תפקידים בעלי כותרת "
              "לא-סטנדרטית או נדירה עלולים להיות מאוחזרים-חסר.")
add_kv_bullet(doc, "שעתוק מבני (K-Means)", "האשכולות משקפים את מבנה השוק הקיים, ועלולים להנציח דפוסים "
              "קיימים (למשל מתאם בין גילוי שכר לסוג חברה) במקום לאתגר אותם.")
add_kv_bullet(doc, "סיכון בזיהוי אנומליות", "סימון משרה כ\"חריגה\" עלול לתייג בטעות מעסיק לגיטימי אך יוצא-דופן "
              "(False Positive). לכן הניסוח הוא אזהרה מייעצת (\"מומלץ לאמת\") ולא פסילה.")
add_kv_bullet(doc, "הטיית מודל השפה (LLM)", "מודלי שפה נושאים הטיות מנתוני האימון שלהם. אנו מצמצמים זאת "
              "בכך שה-LLM אינו מחשב עובדות אלא רק מפענח ומנסח, מעוגן בנתונים דטרמיניסטיים — צמצום, לא ביטול.")

add_lead(doc, "ג. שימוש אחראי ב-AI, שקיפות ופרטיות")
add_bullet(doc, "מניעת הזיות: כל מספר מחושב ב-Pandas וכל דרישה מעוגנת — הסוכן אינו מציג עובדה בדויה כאמת.")
add_bullet(doc, "שקיפות והסברות: ML מדיד וציוני התאמה ניתנים-להסבר, כך שהמשתמש יכול להבין מדוע משרה הומלצה.")
add_bullet(doc, "פרטיות: קורות-חיים מעובדים בזיכרון ואינם נשמרים בצד השרת (גרסאות נשמרות מקומית בדפדפן); "
                "מפתח ה-API מוחרג מבקרת גרסאות; payloads חסומים בגודלם; פלט מטוהר מפני XSS (DOMPurify).")
add_bullet(doc, "גילוי נאות: טקסט קורות-החיים והמשרות נשלח לספק LLM חיצוני (Anthropic) לשם הניסוח — עובדה "
                "שראוי ליידע בה את המשתמש.")

add_lead(doc, "ד. הוגנות ואוטונומיית המשתמש")
add_bullet(doc, "אי-אפליה: ההתאמה מבוססת תפקיד/ניסיון/מיקום/כישורים בלבד, ולא על מאפיינים מוגנים (גיל, מגדר, "
                "מוצא); עם זאת, ייתכנו משתני-תיווך (proxies) ויש להישמר מהם.")
add_bullet(doc, "כנות מול המעסיק: בהתאמת קורות-חיים הסוכן לעולם אינו ממציא ניסיון — הוא מתשאל עובדות אמיתיות "
                "ומנהל פערים ביושר, כדי לא לסייע למשתמש להטעות.")
add_bullet(doc, "כיבוד אוטונומיה: כשמתגלה פער בין שאיפה לרקע, הסוכן נותן \"בדיקת מציאות\" מכבדת אך תמיד מותיר "
                "את ההחלטה למשתמש — הסוכן כלי תומך-החלטה, לא תחליף לשיקול דעת אנושי.")

add_par(doc, "מסמך זה מתעד את כלל ההחלטות שהתקבלו לאורך הפרויקט — מהנדסת הנתונים (אילו קבצים נשמרו, אוחדו "
             "או הוסרו; כיצד נוקתה כל עמודה; ואילו שורות וחריגים סוננו) ועד ארכיטקטורת המודלים והסוכן — "
             "תוך הסבר הנימוק לכל החלטה.", space_before=6)

# ---------- 2. RAW DATA ----------
add_heading(doc, "2. מבנה הנתונים הגולמיים (נקודת הפתיחה)", level=1)
add_par(doc, "מקור הנתונים: מאגר \"LinkedIn Job Postings (2023–2024)\" מ-Kaggle, המכיל 11 קבצי CSV "
             "הקשורים זה לזה. נקודת הפתיחה: 123,849 משרות ו-24,473 חברות. הקבצים מאורגנים סביב שני "
             "מפתחות מרכזיים — job_id (מזהה משרה) ו-company_id (מזהה חברה) — ושני מילוני קודים "
             "(skill_abr, industry_id).")
add_rtl_table(doc,
    ["קובץ", "שורות", "עמודות", "תוכן"],
    [
        ["postings.csv", "123,849", "31", "טבלת העובדות המרכזית — שורה לכל משרה"],
        ["companies.csv", "24,473", "10", "פרופיל חברה (שם, גודל, מיקום)"],
        ["company_industries.csv", "24,375", "2", "חברה ← תעשייה"],
        ["company_specialities.csv", "169,387", "2", "חברה ← תגיות התמחות (טקסט חופשי)"],
        ["employee_counts.csv", "35,787", "4", "תמונות מצב של מספר עובדים/עוקבים לאורך זמן"],
        ["benefits.csv", "67,943", "3", "משרה ← סוג הטבה"],
        ["job_industries.csv", "164,808", "2", "משרה ← industry_id (מקודד)"],
        ["job_skills.csv", "213,768", "2", "משרה ← skill_abr (מקודד)"],
        ["salaries.csv", "40,785", "8", "משרה ← פירוט שכר"],
        ["industries.csv", "422", "2", "מילון: industry_id ← שם תעשייה"],
        ["skills.csv", "35", "2", "מילון: skill_abr ← שם כישור"],
    ],
    col_widths=[1.7, 0.8, 0.8, 3.4])
add_spacer(doc)
add_par(doc, "בעיות מרכזיות שזוהו בשלב הפרופיילינג:", bold=True, space_before=4)
add_bullet(doc, "שכר: קיים רק עבור ~29% מהמשרות, וערכיו משובשים בשל תיוג שגוי של תקופת התשלום "
                "(pay_period) — למשל שכר שנתי שתויג כשעתי והוכפל לערכים של מעל 100 מיליון דולר.")
add_bullet(doc, "גאוגרפיה: ~99.95% מהמשרות בארה\"ב, עם זנב קטן של משרות זרות.")
add_bullet(doc, "כפילויות: אלפי פרסומים חוזרים של אותה משרה.")
add_bullet(doc, "ערכים חסרים ורעש: עמודות עם אחוז ריקים גבוה, ערכי \"0\" מדומים בשדות גאוגרפיים, "
                "תווי קידוד פגומים בטקסט, וכותרות פסולות.")
add_bullet(doc, "סיכון יחס אחד-לרבים (One-to-Many): מספר טבלאות בן מכילות שורות רבות למפתח — סיכון "
                "להתפוצצות שורות במיזוג (ראו פרק 9).")

# ---------- 3. PRINCIPLES ----------
add_heading(doc, "3. עקרונות מנחים", level=1)
add_bullet(doc, "מבנה היעד: טבלה אחת מאוחדת (Master Table), שורה אחת לכל משרה (job_id).")
add_bullet(doc, "שימור נתונים: ככלל לא מוחקים שורות בשל ערכים חסרים — חוסר מקודד כדגל/קטגוריה — "
                "כדי לשמר את היקף ה-Big Data ולמנוע הטיה באשכול. חריגים מוסרים רק מסיבות איכות מובהקות.")
add_bullet(doc, "מסד נתונים לארה\"ב בלבד (USA-only): הסינון מתבצע לפי מיקום המשרה (לא מטה החברה).")
add_bullet(doc, "כל עמודה מקובצת/מקודדת מקבלת מילון שמור (data_dictionary.json).")
add_bullet(doc, "התאמה למודלים: הסכמה תוכננה כך שתזין ישירות את מודלי הדמיון והאשכול.")

# ---------- 4. FILE LEVEL ----------
add_heading(doc, "4. החלטות ברמת הקבצים (שמירה / איחוד / הסרה)", level=1)
add_rtl_table(doc,
    ["קובץ", "החלטה", "נימוק מרכזי"],
    [
        ["postings.csv", "שמירה (Keep)", "טבלת העובדות המרכזית; כל הקבצים מתחברים אליה"],
        ["companies.csv", "שמירה (Keep)", "מימד חברות נקי, שורה לכל חברה; מקור שם החברה"],
        ["company_industries.csv", "איחוד → company_industry", "טקסט קריא, כיסוי 98.5%, חיבור ~1:1"],
        ["company_specialities.csv", "הסרה (Drop)", "82,960 תגיות טקסט חופשי לא עקביות, עד 84 לחברה, 27% חסרים — רעש"],
        ["employee_counts.csv", "איחוד → ממלא company_size ואז מוסר", "ריבוי תמונות מצב (עד 13/חברה); נשמר העדכני בלבד"],
        ["job_skills.csv", "איחוד → skills", "תרגום קודים→שמות, רשימה אחת למשרה; ליבת התאמת Role"],
        ["job_industries.csv", "איחוד → job_industry", "תרגום קודים→שמות, ערך ראשון תקף למשרה"],
        ["benefits.csv", "הסרה (Drop)", "כיסוי 23% בלבד, 60% מנוחשים אוטומטית, אינו פילטר ליבה"],
        ["salaries.csv", "הסרה (Drop)", "יתירות מלאה עם שדות השכר ב-postings; שורות נוספות הן יתומות"],
        ["industries.csv", "מילון (משמש ומוסר)", "מתרגם industry_id ← שם"],
        ["skills.csv", "מילון (משמש ומוסר)", "מתרגם skill_abr ← שם"],
    ],
    col_widths=[1.9, 1.9, 2.9])

# ---------- 5. COLUMN LEVEL ----------
add_heading(doc, "5. החלטות ברמת העמודות", level=1)
add_heading(doc, "5.1 עמודות שנוקו / שונו / נוצרו", level=2)
add_kv_bullet(doc, "title", "ניקוי רווחים ושורות חדשות.")
add_kv_bullet(doc, "experience_level", "ריקים → \"Not Specified\"; נוספה עמודה מספרית סדורה experience_rank (0–6).")
add_kv_bullet(doc, "remote_allowed", "ריק → 0 (דגל בינארי; 1 = מאפשר עבודה מרחוק).")
add_kv_bullet(doc, "location", "פוצל ל-job_city ו-job_state (≈85% נפרסים); המקור נשמר.")
add_kv_bullet(doc, "salary", "נורמליזציה לשכר שנתי ב-USD + תיקון חכם של תיוג pay_period (ראו פרק 7); נוספה salary_band מספרית (0–7).")
add_kv_bullet(doc, "skills", "תרגום קודים→שמות, רשימה מופרדת בפסיקים; נוספה skill_count (0–3).")
add_kv_bullet(doc, "job_industry", "תרגום, ערך יחיד ראשון תקף; נשמרו 374 קטגוריות מפורטות.")
add_kv_bullet(doc, "company_size", "סולם 0–7; השלמת חסרים מ-employee_count; 0 = לא ידוע.")
add_kv_bullet(doc, "company_name", "ניקוי רווחים; כל הערכים נשמרו (כולל \"Confidential\" ושמות קצרים לגיטימיים כמו EY/HP/3M).")
add_kv_bullet(doc, "posting_date", "המרת listed_time מ-epoch-ms לתאריך קריא.")
add_kv_bullet(doc, "text_blob", "שדה ML מנוקה (title+skills+description): אותיות קטנות, הסרת URL/אימייל/HTML/קידוד פגום, סינון מילות עצירה.")
add_kv_bullet(doc, "job_posting_url, application_url", "נשמרו מתוך postings כעמודות מטא-דאטה אקטיביות — כדי להעצים את הסוכן ביכולת ביצוע (Action-Oriented Agent): ניתוב המשתמש ישירות לעמוד המשרה ולמסלול הגשת המועמדות. אינן משמשות כמאפייני ML ומטופלות כמטא-דאטה פסיבי בלבד (אינן נכנסות ל-TF-IDF או ל-K-Means).")

add_heading(doc, "5.2 עמודות שהוסרו", level=2)
add_par(doc, "company_id (מפתח חיבור — מוסר לאחר המיזוג), work_type (כפילות CAPS), views, applies, "
             "currency (הכול USD לאחר סינון לארה\"ב), שדות שכר גולמיים (min/max/med, pay_period, normalized), "
             "compensation_type ו-sponsored (ערך יחיד), skills_desc ו-closed_time (~99% ריקים), "
             "country/state/city של מטה החברה (מיותר מול מיקום המשרה), application_type, "
             "zip_code, fips, וכן has_salary/geo_granularity/company_size_unknown (הוחלפו בקידודים יעילים יותר). "
             "הערה: קישורי job_posting_url ו-application_url, שהוסרו בגרסה מוקדמת, הוחזרו למאגר (ראו סעיף 5.1) "
             "כדי לאפשר לסוכן יכולת פעולה והפניה ישירה להגשת מועמדות.")

# ---------- 6. ROWS / OUTLIERS ----------
add_heading(doc, "6. סינון שורות וחריגים (Outliers)", level=1)
add_par(doc, "סה\"כ הוסרו 4,732 שורות — מ-123,849 ל-119,117 משרות (עדיין מעל סף ה-Big Data של 100K).", bold=True)
add_kv_bullet(doc, "3 שורות — כותרות פסולות", "הכותרות (\"200052\", \"![\") חסרות שדה Role שמיש.")
add_kv_bullet(doc, "4,673 שורות — פרסומים כפולים", "אותה משרה שפורסמה שוב (title + company + location + description זהים; שכר זהה ב-98.8%). נשמר העותק עם מספר הצפיות הגבוה ביותר.")
add_kv_bullet(doc, "12 שורות — משרות זרות", "מיקום מחוץ לארה\"ב (קנדה, פיליפינים, הולנד ועוד) — מסד לארה\"ב בלבד.")
add_kv_bullet(doc, "43 שורות — שכר מתחת למינימום במשרה מלאה", "שכר שנתי < $15,080 במשרה מלאה — אינדיקציה לנתון שגוי/מודעה לא אמינה (למשל \"רופא\" ב-$10,000).")

# ---------- 7. SALARY ----------
add_heading(doc, "7. טיפול חכם בחריגי שכר (תיקון במקום מחיקה)", level=1)
add_par(doc, "שורש הבעיה: תיוג שגוי של pay_period — שכר שנתי שתויג כשעתי/חודשי וכדומה, וכן ערכים שהוזנו "
             "ביחידות של אלפים. במקום למחוק את הערכים החריגים, זוהה הגורם ותוקן, כך ש-~420 ערכי שכר אמיתיים "
             "שוחזרו ורק חריגים שאינם ניתנים לשחזור רוקנו. מקדמי הנרמול: שעתי×2080, שבועי×52, דו-שבועי×26, חודשי×12, שנתי×1.")
add_par(doc, "כללי השחזור (מיושמים לפי הסדר):", bold=True, space_before=4)
add_bullet(doc, "ערך מנורמל > $500K אך הבסיס הוא שכר שנתי סביר ($10K–$500K) → שימוש בבסיס (היה שנתי).", sub=True)
add_bullet(doc, "YEARLY עם בסיס < 50 → שכר שעתי → ×2080.", sub=True)
add_bullet(doc, "YEARLY עם בסיס 50–999 → הוזן באלפים → ×1000 (למשל \"Accounting Manager 90–120\" = $90K–$120K).", sub=True)
add_bullet(doc, "MONTHLY עם בסיס < 500 → תיוג שגוי של שעתי → ×2080.", sub=True)
add_bullet(doc, "HOURLY עם max פגום (max>500 ו-min≤250) → שימוש ב-min התקין × 2080.", sub=True)
add_bullet(doc, "ערכי 0/1 (ללא תשלום/placeholder) → רוקנו (השורה נשמרה).", sub=True)
add_bullet(doc, "ערכים מעל $1M לאחר השחזור → נשמרו (משתכרים גבוהים לגיטימיים).", sub=True)
add_bullet(doc, "ערכים שאינם ניתנים לשחזור במשרה מלאה (מתחת לשכר המינימום) → השורות הוסרו (ראו פרק 6).", sub=True)
add_par(doc, "טיפולים נוספים:", bold=True, space_before=4)
add_bullet(doc, "15 משרות בארה\"ב במטבע לא-USD → תוקנו ל-USD (למשל \"England, AR\" שתויג GBP).")
add_bullet(doc, "שכר נמוך אמיתי במשרות חלקיות/חוזה (Part-time/Contract) → נשמר.")
add_bullet(doc, "שכר חסר (70.4%) → קודד כ-salary_band = 0 (\"לא צוין\"), ללא השלמה מלאכותית.")

# ---------- 8. FINAL SCHEMA + DICTS ----------
add_heading(doc, "8. הסכמה הסופית והמילונים", level=1)
add_par(doc, "מאגר היעד כולל 22 עמודות (119,117 שורות), בשתי שכבות: עמודות קריאות לסוכן (כולל קישורי "
             "פעולה לעמוד המשרה ולהגשת מועמדות) + מאפיינים מספריים/דגלים למודלים + שדה text_blob לדמיון.", space_after=4)
add_par(doc, "job_id, title, description, location, job_posting_url, application_url, job_city, job_state, "
             "posting_date, work_type, experience_level, remote_allowed, salary, salary_band, skills, "
             "skill_count, job_industry, company_name, company_industry, company_size, experience_rank, text_blob",
        bold=True, font=HEAD_FONT, size=10)

add_heading(doc, "8.1 מילונים (לכל עמודה מקודדת)", level=2)
add_par(doc, "salary_band:", bold=True)
add_rtl_table(doc, ["קוד", "טווח"],
    [["0", "לא צוין"], ["1", "מתחת ל-$40,000"], ["2", "$40,000–$60,000"], ["3", "$60,000–$80,000"],
     ["4", "$80,000–$100,000"], ["5", "$100,000–$150,000"], ["6", "$150,000–$200,000"], ["7", "$200,000 ומעלה"]],
    col_widths=[0.7, 2.5])
add_spacer(doc)
add_par(doc, "company_size:", bold=True)
add_rtl_table(doc, ["קוד", "מספר עובדים"],
    [["0", "לא ידוע"], ["1", "1–10"], ["2", "11–50"], ["3", "51–200"],
     ["4", "201–500"], ["5", "501–1,000"], ["6", "1,001–5,000"], ["7", "5,001+"]],
    col_widths=[0.7, 2.5])
add_spacer(doc)
add_par(doc, "experience_rank:", bold=True)
add_rtl_table(doc, ["קוד", "רמת ניסיון"],
    [["0", "לא צוין"], ["1", "Internship"], ["2", "Entry level"], ["3", "Associate"],
     ["4", "Mid-Senior level"], ["5", "Director"], ["6", "Executive"]],
    col_widths=[0.7, 2.5])

# ---------- 9. MERGE SAFETY ----------
add_heading(doc, "9. בטיחות המיזוג (מניעת התפוצצות שורות)", level=1)
add_par(doc, "ארבע טבלאות הן ביחס אחד-לרבים, ומיזוג ישיר שלהן היה מנפח את הטבלה ליצירת שורות כפולות שגויות: "
             "job_skills (×3), job_industries (×3), employee_counts (×13), company_industries (×2).")
add_par(doc, "הפתרון: כל טבלת אחד-לרבים מצומצמת לשורה אחת למפתח לפני המיזוג, ולאחר מכן מתבצע left-join אל "
             "הטבלה המרכזית. בקוד נעשה שימוש ב-merge(validate=...) של pandas ובבדיקת מספר שורות לאחר כל מיזוג, "
             "כך שכל ניפוח מקרי יעצור את התהליך בשגיאה. שני מילוני התרגום (skills, industries) הם 1:1 ובטוחים.")

# ---------- 10. ML ARCHITECTURE CHAPTER ----------
add_heading(doc, "פרק 10: ארכיטקטורת מודלי למידת המכונה, זיהוי אנומליות ומערכת ההמלצות האקטיבית", level=1)
add_par(doc, "לאחר בניית ה-Master Table, נבנתה שכבת בינה המשלבת שלושה מודלים שנלמדו בכיתה, על מנת להפוך את "
             "המאגר לסוכן המלצות אקטיבי. הארכיטקטורה היברידית ומורכבת משלושה רכיבים משלימים: (1) סינון גאוגרפי "
             "קשיח (Hard Filtering) לפי מיקום המשרה; (2) דמיון סמנטי (Cosine Similarity) להתאמת תפקיד מתוך שאילתה "
             "בשפה טבעית; (3) אשכול פרופילים (K-Means) לקיבוץ המשרות לפי ניסיון, שכר וגודל מעסיק, המשמש גם לזיהוי "
             "אנומליות בשוק.")

add_heading(doc, "10.1 אשכול פרופילי משרה (K-Means Clustering)", level=2)
add_par(doc, "מטרת המודל: קיבוץ המשרות לקבוצות הומוגניות (\"פרופילי שוק\") וזיהוי דפוסים נסתרים. נבחרו ארבעה "
             "מאפיינים מספריים/סדורים מקודדים: experience_rank, company_size, salary_band, remote_allowed. "
             "המאפיינים עברו תקנון (StandardScaler) כדי לאזן את השפעת הסקאלות השונות על מדד המרחק.")
add_par(doc, "בחירת מספר האשכולות (K) בוצעה בשיטת המרפק (Elbow Method): חושב ה-Inertia (סכום ריבועי המרחקים "
             "התוך-אשכוליים) עבור K בין 2 ל-10, ונבחר באופן אובייקטיבי ה-K בעל המרחק המרבי מהקו הישר שבין "
             "הקצוות. התוצאה: K=5. האשכולות שהתקבלו מאוזנים (ללא אשכול דומיננטי), ומפרשים מגזרי שוק ברורים:")
add_rtl_table(doc, ["אשכול", "נתח", "פרופיל (פרשנות שוק)"],
    [["3", "31.8%", "משרות לכניסה לשוק בחברות גדולות מאוד; שכר לא מפורסם"],
     ["0", "25.5%", "משרות בכירות-בינוניות בחברות גדולות; שכר לא מפורסם"],
     ["1", "16.4%", "משרות לכניסה לשוק בחברות קטנות"],
     ["2", "14.2%", "משרות בכירות-בינוניות עם שכר מפורסם (כ-$80K–$100K)"],
     ["4", "12.2%", "משרות מרחוק (100% Remote) — מגזר נבדל"]],
    col_widths=[0.7, 0.9, 4.0])

add_heading(doc, "10.2 דמיון קוסינוס (Cosine Similarity) — התאמת שפה טבעית", level=2)
add_par(doc, "להתאמת התפקיד מתוך תיאור חופשי של המשתמש, הומר השדה text_blob (title+skills+description לאחר ניקוי) "
             "לייצוג וקטורי באמצעות TF-IDF (Term Frequency – Inverse Document Frequency), המשקלל מילים לפי "
             "תדירותן במשרה ביחס לנדירותן בכלל המאגר. הדמיון בין שאילתת המשתמש למשרה מחושב כקוסינוס הזווית בין "
             "הווקטורים:")
add_par(doc, "cos(q, d) = (q · d) / ( ||q|| × ||d|| )", bold=True, font=HEAD_FONT, size=11)
add_par(doc, "ערך קרוב ל-1 מציין התאמה סמנטית גבוהה. גישה זו לוכדת מונחים ספציפיים (כגון \"python\", \"sql\") "
             "המופיעים בתיאורי המשרה, גם כאשר אינם נמנים עם 35 קטגוריות הכישורים.")

add_heading(doc, "10.3 דמיון ז'קארד (Jaccard Similarity) — חיתוך קבוצות כישורים", level=2)
add_par(doc, "להתאמת כישורים, חושב מדד ז'קארד בין קבוצת הכישורים שהזין המשתמש (A) לבין קבוצת קטגוריות הכישורים "
             "של המשרה (B). המדד הוא יחס גודל החיתוך לגודל האיחוד:")
add_par(doc, "J(A, B) = |A ∩ B| / |A ∪ B|", bold=True, font=HEAD_FONT, size=11)
add_par(doc, "ערך 1 מציין חפיפה מלאה, 0 מציין היעדר חפיפה. המדד מתאים במיוחד לקבוצות בדידות כמו תגיות כישורים.")

add_heading(doc, "10.4 זיהוי אנומליות — מרחק מתמטי מהצנטרואיד", level=2)
add_par(doc, "לכל משרה חושב המרחק האוקלידי בין וקטור המאפיינים המתוקנן שלה לבין צנטרואיד האשכול שאליו שויכה "
             "(d = √Σ(xᵢ − cᵢ)²). משרה הרחוקה מאוד ממרכז הקבוצה שלה מייצגת צירוף מאפיינים חריג בשוק "
             "(לדוגמה: שכר גבוה במיוחד לצד דרישת ניסיון אפסית). סף האנומליה הוגדר כאחוזון ה-95 של התפלגות "
             "המרחקים; כל משרה מעבר לסף סומנה בדגל is_anomaly=1.")
add_par(doc, "תוצאה: סף = 2.13 (אחוזון 95), 5,931 משרות (5.0%) סומנו כאנומליות. התפלגות המרחקים וקו הסף "
             "מוצגים בהיסטוגרמה ייעודית בקוד.")

add_heading(doc, "10.5 הסוכן האקטיבי — ציון משולב, התרעות סיכון וניתוב לפעולה", level=2)
add_par(doc, "פונקציית get_recommendations מממשת את הזרימה ההיברידית: תחילה סינון גאוגרפי קשיח (לפי job_state/"
             "job_city), ולאחריו דירוג לפי ציון משולב — 50% דמיון קוסינוס (תפקיד) ועוד 50% דמיון ז'קארד (כישורים):")
add_par(doc, "Score = 0.5 × Cosine + 0.5 × Jaccard", bold=True, font=HEAD_FONT, size=11)
add_bullet(doc, "פעולה אקטיבית (Action-Oriented): עבור כל המלצה מוצגים job_posting_url ו-application_url, כדי "
                "שהסוכן ינתב את המשתמש ישירות לעמוד המשרה ולהגשת המועמדות.")
add_bullet(doc, "פרופיילינג סיכונים (Risk Alert): אם משרה מומלצת מסומנת is_anomaly=1, הסוכן מצרף אזהרה יזומה: "
                "\"AGENT NOTICE: משרה זו סומנה כאנומליה בשוק בשל צירוף מאפיינים חריג — מומלץ לאמת את הנתונים\", "
                "כדי להגן על המשתמש מפני מודעות חשודות או שגויות.")

# ---------- 11. NATURAL-LANGUAGE / TITLE+EXPERIENCE RAG CHAPTER ----------
add_heading(doc, "פרק 11: ממשק שפה טבעית מלא וארכיטקטורת סוכן RAG מבוססת ניסיון וכותרת", level=1)
add_par(doc, "בגרסת הסוכן הסופית עברנו לממשק שפה טבעית מלא (End-to-End Natural Language): המשתמש מקליד "
             "בקשה חופשית אחת (לדוגמה: \"I am looking for a Senior Data Analyst position in NY\"), והסוכן "
             "מפענח אותה, מאחזר מקומית את המשרות המתאימות ביותר, ומנסח תשובה אישית בעברית. במעבר זה "
             "הוצא משימוש (deprecated) אינדקס הכישורים המבני שהתבסס על דמיון Jaccard, לטובת צינור שפה "
             "טבעית מקצה לקצה.")

add_heading(doc, "11.1 פענוח הבקשה וסינון גאוגרפי קשיח", level=2)
add_par(doc, "הבקשה החופשית מנותחת תכנותית: אם מופיע בה קוד מדינה (כגון NY, CA), מוחל סינון גאוגרפי קשיח "
             "על העמודה job_state, כך שהדירוג מתבצע רק על משרות במדינה המבוקשת. בהיעדר קוד מדינה — החיפוש "
             "מתבצע על כל ארה\"ב.")

add_heading(doc, "11.2 התאמה דו-רכיבית מבוססת Cosine (כותרת + רמת ניסיון)", level=2)
add_par(doc, "במקום חיתוך קבוצות כישורים (Jaccard), מחושבים שני וקטורי דמיון קוסינוס נפרדים מעל ייצוג ה-TF-IDF "
             "(Scikit-Learn / NumPy, בזיכרון):")
add_bullet(doc, "ציון כותרת (Title Score): דמיון קוסינוס בין בקשת המשתמש לבין עמודת ה-title של המשרות.")
add_bullet(doc, "ציון ניסיון (Experience Score): דמיון קוסינוס בין הבקשה לבין עמודת experience_level.")
add_par(doc, "השניים מאוחדים לציון היברידי: Score = 0.6 × Title + 0.4 × Experience. הכותרת מקבלת משקל "
             "גבוה יותר משום שהיא המבע הישיר ביותר של התפקיד המבוקש, בעוד רמת הניסיון מעדנת את הדירוג. "
             "נבחרות שלוש המשרות בעלות הציון הגבוה ביותר.", space_before=2)

add_heading(doc, "11.3 סינתזה לשונית ופרופיל סיכון", level=2)
add_par(doc, "עובדות שלוש המשרות (כותרת, חברה, מיקום, סרגל שכר, רמת ניסיון, דגל is_anomaly וקישורי הפעולה) "
             "מועברות אל Claude (Anthropic API, מודל Sonnet 4.6 — גרסת 3.5 הוצאה משימוש), המנסח תשובה "
             "מותאמת ומעודדת בעברית המסבירה מדוע כל משרה תואמת את הכותרת ורמת הניסיון שביקש המשתמש.")
add_bullet(doc, "פרופיל סיכון: עבור משרה שסומנה is_anomaly=1, הסוכן מדגיש אזהרה (\"⚠️ התראת סיכון סוכן\") "
                "ומסביר בשפה פיננסית-מקצועית כי הפרסום סוטה סטטיסטית מקו הבסיס של אשכול השוק, ומומלץ לאמתו.")
add_bullet(doc, "פעולה: מוצגים קישורי job_posting_url ו-application_url כ-Markdown פעיל, לפעולה מיידית.")
add_bullet(doc, "כל חישובי הדמיון מתבצעים מקומית בזיכרון (In-Memory RAG) ללא מסד נתונים וקטורי חיצוני; "
                "מפתח ה-API נטען מקובץ local.env המוחרג מבקרת גרסאות.")

# ---------- פרק 12: DATA-GROUNDED CHATBOT ----------
add_heading(doc, "פרק 12: צ'אטבוט תעסוקתי מבוסס-נתונים — נתב כוונות ומנוע אנליטיקה דטרמיניסטי", level=1)
add_par(doc, "בשלב זה הורחב הסוכן ממנוע חיפוש בלבד לעוזר קריירה שיחתי. מעבר לאיתור משרות, המערכת מסוגלת כעת "
             "לענות על שאלות ייעוץ פתוחות (הכנה לריאיון, שיפור קו\"ח, אילו כישורים ללמוד, ציפיות שכר) וכן "
             "לחשב תשובות מבוססות-נתונים אמיתיות מעל מאגר ה-Gold (כ-119 אלף משרות). כל הודעת משתמש מנותבת "
             "אוטומטית להתנהגות המתאימה, דרך אותו תיבת קלט אחת.")

add_heading(doc, "12.1 נתב הכוונות (Intent Router)", level=2)
add_par(doc, "דגל החיפוש הבינארי הקודם (is_search) הוחלף בשדה כוונה רב-ערכי, mode, המופק על-ידי אותה קריאת "
             "פענוח קיימת (parse_query) — כך שחיפוש רגיל אינו עולה קריאת-מודל נוספת. ערכי הכוונה:")
add_bullet(doc, "search — איתור ודירוג משרות (כבעבר).")
add_bullet(doc, "data_question — שאלה על נתון מחושב מעל אוכלוסיית משרות (ממוצע, ספירה, הנפוצים ביותר).")
add_bullet(doc, "advice — ייעוץ קריירה פתוח (ריאיון, קו\"ח, כישורים, שכר).")
add_bullet(doc, "smalltalk — פתיח / תודה / \"מה אתה יכול לעשות\".")
add_bullet(doc, "more — בקשה לעוד תוצאות של אותו חיפוש קודם (Pagination).")
add_par(doc, "במקרה של עמימות בין search ל-data_question, ברירת המחדל היא search — כך שניתוב שגוי לעולם אינו "
             "מדרדר את החוויה מתחת לקיים. זוהי שכבת החלטה יחידה, שקופה ומוסברת.", space_before=2)

add_heading(doc, "12.2 מנוע האנליטיקה הדטרמיניסטי (analytics.py)", level=2)
add_par(doc, "ליבת התרומה מבוססת-הנתונים. מימוש כ\"רשימת היתר\" (Whitelist) סגורה של שבע פונקציות צבירה (Aggregation) "
             "מעל מסגרת הנתונים — המספר מחושב תמיד ב-Pandas, ולעולם לא על-ידי המודל. צינור בן שלושה שלבים:")
add_bullet(doc, "מיפוי (Map): Claude בוחר פונקציה אחת ופרמטרים מתוך הרשימה הסגורה ומחזיר JSON — בחירה דרך "
                "מילון (dict lookup) בלבד, ללא eval / getattr / exec, כך שלא ניתן להריץ קוד שרירותי.")
add_bullet(doc, "חישוב (Compute): Python מריץ את אותה פונקציה ומחזיר עובדה מספרית (תמיד כולל n = גודל האוכלוסייה).")
add_bullet(doc, "ניסוח (Narrate): Claude מנסח את העובדה במשפט עברי קצר אחד, תוך שימוש אך ורק במספרים שחושבו.")
add_par(doc, "שבע הפונקציות:", bold=True, space_before=2)
add_rtl_table(doc,
    ["פונקציה", "על מה היא עונה", "חישוב"],
    [
        ["salary_stats", "השכר הטיפוסי לתפקיד/מיקום", "חציון ואחוזונים (p25/p75) של עמודת השכר הרציפה"],
        ["top_skills", "אילו כישורים נדרשים", "פיצול רשימת הכישורים וספירת תדירויות"],
        ["count_jobs", "כמה משרות עונות לתנאים", "ספירת שורות לאחר סינון"],
        ["remote_share", "איזה חלק מהמשרות מרחוק", "ממוצע remote_allowed (לאחר השמטת ערכים חסרים)"],
        ["experience_breakdown", "התפלגות רמות הניסיון", "ספירת ערכים על experience_level"],
        ["top_locations", "היכן מרבית המשרות", "ספירת ערכים על job_state"],
        ["top_industries", "אילו תעשיות מגייסות", "ספירת ערכים על job_industry"],
    ])
add_par(doc, "סינון התפקיד (role) אינו השוואת מחרוזות נאיבית אלא שימוש חוזר בדמיון הקוסינוס מעל ייצוג ה-TF-IDF "
             "של הכותרות (סף 0.30) — כך שה-ML שכבר נבנה הוא שמבצע את בחירת השורות. כל פילטר (תפקיד / מדינה / "
             "ניסיון / מרחוק) אופציונלי וניתן לשילוב חופשי (שער AND).", space_before=2)
add_par(doc, "יושרה אקדמית: רשימת פונקציות סגורה (ללא יצירת קוד על-ידי המודל) מבטיחה שחזוריות, בטיחות ויכולת "
             "הסבר; ובמקרה של n קטן (מתחת ל-5) הניסוח מציין במפורש שהתשובה מבוססת על מעט פרסומים בלבד.")

add_heading(doc, "12.3 ייעוץ מעוגן ושילוב בזרם קורות-החיים והריאיון", level=2)
add_par(doc, "מצב ה-advice מנסח תשובת ייעוץ עברית קצרה, המעוגנת בהקשר: שיחת ההיסטוריה האחרונה, המשרות "
             "המוצגות כעת על המסך, ובמידת הצורך גם פרופיל קורות-החיים — כך שתשובה לשאלה \"האם המשרה הראשונה "
             "מתאימה לי?\" מתייחסת למשרה האמיתית, לא להשערה.")
add_par(doc, "שתי היכולות החדשות (אנליטיקה וייעוץ) הן שירותים משותפים, הנגישים משני מסלולי הכניסה: גם מהנתב "
             "ללא קו\"ח, וגם ממסלול קורות-החיים/הריאיון (שבסיומו נטען פרופיל המשתמש) — דרך הוספת פעולות "
             "data_question ו-advice לבקר השיחה (cv/converse.py). בלשונית הסריקה החיה (ישראל), שבה כיסוי "
             "השכר והתעשייה דליל, שאלת data_question מנותבת בכוונה לייעוץ במקום לנתון מחושב.", space_before=2)

# ---------- 13. DELIVERABLES ----------
add_heading(doc, "13. תוצרים", level=1)
add_bullet(doc, "analysis_and_profiling.py — פרופיילינג נתונים גולמיים והדמיית חריגים (1.5×IQR).")
add_bullet(doc, "preperation_and_merge.py — בניית ה-Master Table: ניקוי, שחזור שכר, סינון כפילויות/חריגים, "
                "איחוד טבלאות, text_blob, ייצוא data_dictionary.json, ומיזוגים עם validate='m:1'.")
add_bullet(doc, "training_and_clustering.py — אשכול K-Means, זיהוי אנומליות, דמיון קוסינוס וז'קארד, "
                "ומערכת ההמלצות האקטיבית.")
add_bullet(doc, "agent_runner.py — סוכן ה-RAG המקומי: שליפה מקומית + סינתזה לשונית בעברית באמצעות Claude.")
add_bullet(doc, "analytics.py — מנוע האנליטיקה הדטרמיניסטי: שבע פונקציות צבירה ב\"רשימת היתר\" וצינור "
                "מיפוי→חישוב→ניסוח לתשובות מבוססות-נתונים (פרק 12).")
add_bullet(doc, "פלטים: master_jobs_dataset.csv (מלא, UTF-8) + גרסה מצומצמת, data_dictionary.json, "
                "gold_linkedin_with_clusters.csv (כולל cluster_id ו-is_anomaly), ומודלים שמורים (joblib).")
add_bullet(doc, "מסמך תיעוד זה (.docx) בעברית, המנמק כל החלטה.")

# ---------- 14. CONCLUSIONS & APPLICATION ----------
add_heading(doc, "14. מסקנות ויישום", level=1)
add_par(doc, "פרק זה מסכם את מסקנות הפרויקט, דן ביישום שלו בתרחישים בעולם האמיתי (הסבר בלבד), ובהנחות "
             "והמגבלות שלקחנו בחשבון.")

add_lead(doc, "14.1 מסקנות עיקריות")
add_bullet(doc, "הארכיטקטורה ההיברידית הוכיחה את עצמה: שילוב ML קלאסי ומדיד עם LLM מניב סוכן שקוף ועמיד-"
                "להזיות — תזת הליבה של הפרויקט אוששה.")
add_bullet(doc, "הנדסת הנתונים היא היסוד: עיקר המאמץ והערך היו בהפיכת 11 טבלאות רלציוניות רועשות לטבלת אב "
                "אחת אמינה; איכות ההמלצות תלויה ישירות בכך.")
add_bullet(doc, "עיגון גובר על LLM גולמי: הגבלת המודל לפענוח וניסוח מעל עובדות דטרמיניסטיות נתנה תשובות "
                "אמינות ושחזוריות.")
add_bullet(doc, "K-Means וזיהוי האנומליות חשפו מבנה שוק אמיתי וסיפקו אות סיכון שימושי למשתמש.")

add_lead(doc, "14.2 יישום בעולם האמיתי (תרחישים — הסבר בלבד)")
add_kv_bullet(doc, "מחפש עבודה", "עוזר קריירה אישי: חיפוש בשפה טבעית לפי תפקיד/ניסיון/מיקום, התרעות סיכון "
              "על מודעות חריגות, והתאמת קורות-חיים למשרה.")
add_kv_bullet(doc, "בוגרים ומסיימי לימודים", "בניית קורות-חיים מאפס והצעת תפקידים מתאימים מדורגים לפי אחוז "
              "התאמה (ML).")
add_kv_bullet(doc, "יועצי קריירה והשמה", "לוח-מחוונים תומך-החלטה: שכר טיפוסי, כישורים מבוקשים והיכן מרוכזות "
              "המשרות — מבוסס נתונים אמיתיים.")
add_kv_bullet(doc, "מעסיקים ומגייסים", "השוואת מודעה למבנה השוק (אשכול/אנומליה) כדי לכייל ציפיות שכר "
              "ודרישות.")
add_kv_bullet(doc, "הרחבה גאוגרפית", "הסורק החי (ישראל) מדגים שהארכיטקטורה ניתנת להכללה לשווקים נוספים בהינתן "
              "מקור נתונים.")

add_lead(doc, "14.3 הנחות שלקחנו (Assumptions)")
add_bullet(doc, "מסד לארה\"ב בלבד: סינון לפי מיקום המשרה, המרת כל השכר ל-USD, והנחה שדומיננטיות ארה\"ב "
                "במאגר מצדיקה הסרת הזנב הזר.")
add_bullet(doc, "גרעין שורה-אחת-למשרה (job_id) מספיק; תעשייה ראשית אחת למשרה (81% ממילא יחידה) — כמעט ללא "
                "אובדן מידע.")
add_bullet(doc, "מודעות LinkedIn הן פרוקסי סביר לשוק העבודה (בעיקר צווארון-לבן), והשדות (כותרת, רמת ניסיון) "
                "מדויקים דיים לדירוג.")
add_bullet(doc, "שיבושי השכר עוקבים אחר דפוסים ניתנים-לזיהוי (כללי השחזור), והנחת משרה-מלאה לצורך סף שכר "
                "המינימום.")
add_bullet(doc, "הבקשה החופשית של המשתמש מבטאת את כוונתו האמיתית; עברית ואנגלית מטופלות דרך פענוח ה-LLM.")

add_lead(doc, "14.4 מגבלות (Limitations)")
add_bullet(doc, "כיסוי והטיה: מאגר ממוקד-ארה\"ב וממוקד-LinkedIn, עם שכר דליל (~70% חסר) — ראו פרק 1.10.")
add_bullet(doc, "סיווג כוונה בקריאה אחת עלול לטעות בבקשות עמומות (ראו 1.9).")
add_bullet(doc, "הכישורים הם קטגוריות LinkedIn רחבות ולא גרגריות — התאמה עדינה נשענת על טקסט התיאור בלבד.")
add_bullet(doc, "תמונת מצב סטטית (למעט לשונית הסורק החי), אינה משקפת את השוק בזמן אמת.")
add_bullet(doc, "תלות ב-LLM: דורש מפתח API, רשת ועלות/השהיה; ללא מפתח המערכת מתפקדת בדרגרדציה מבוקרת.")
add_bullet(doc, "אנומליה = חריג סטטיסטי, לא הונאה מאומתת — מייעצת בלבד.")
add_bullet(doc, "הדוגמאות במסמך הן ייצוגיות וממחישות התנהגות, ואינן הערכה חיה ממצה.")

add_lead(doc, "14.5 עבודה עתידית")
add_par(doc, "embeddings צפופים לשדרוג סמנטי; שלב אישור/הבהרה כאשר ביטחון סיווג הכוונה נמוך; רענון נתונים "
             "חי; הרחבה לשווקים גאוגרפיים נוספים; ביקורת הוגנות על משתני-תיווך; והערכה כמותית מבוססת פלטים "
             "אמיתיים שנלכדו.")

# ---------- 15. IMPACT ASSESSMENT ----------
add_heading(doc, "15. הערכת השפעה", level=1)
add_par(doc, "פרק זה מעריך את ההשפעה הפוטנציאלית של הפתרון על תעשיית הגיוס והתעסוקה ועל המשתמשים בה — "
             "היתרונות מצד אחד, והאתגרים והסיכונים מצד שני.")

add_lead(doc, "15.1 יתרונות והשפעה חיובית")
add_kv_bullet(doc, "הנגשת ייעוץ קריירה", "ממשק שפה טבעית מוריד את חסם הכניסה — גם מי שאין לו יועץ קריירה "
              "מקבל הכוונה אישית ומבוססת-נתונים.")
add_kv_bullet(doc, "יעילות למחפש העבודה", "קיצור משמעותי של זמן החיפוש, צמצום עומס המידע, והצפת המשרות "
              "הרלוונטיות באמת.")
add_kv_bullet(doc, "שקיפות בתעשייה אטומה", "התאמה ניתנת-להסבר, ניהול פערים כן והתרעות סיכון מסייעים למשתמש "
              "לקבל החלטות מושכלות ולהיזהר ממודעות חשודות.")
add_kv_bullet(doc, "אינטליגנציית שוק", "תובנות מבוססות-נתונים (שכר טיפוסי, כישורים מבוקשים) מעצימות את "
              "המשתמש במשא-ומתן ובבחירת מסלולי הכשרה.")
add_kv_bullet(doc, "מדרגיות ועלות נמוכה", "אחזור בזיכרון ועלות LLM זניחה מאפשרים פריסה רחבה ונגישה.")

add_lead(doc, "15.2 אתגרים וסיכונים")
add_kv_bullet(doc, "הגברת הטיות בקנה מידה", "פריסה רחבה עלולה להפוך הטיות נתונים/מודל (ארה\"ב, LinkedIn, "
              "צווארון-לבן) להדרה שיטתית של אוכלוסיות מיוצגות-חסר.")
add_kv_bullet(doc, "הומוגניזציה ומרוץ מילות-מפתח", "אם כולם מכווננים קורות-חיים לאותם אותות ML, עלול להיווצר "
              "מרוץ אופטימיזציה שמדכא גיוון ומקוריות.")
add_kv_bullet(doc, "תלות-יתר והטיית אוטומציה", "משתמשים עלולים להאציל שיקול דעת לסוכן; חשוב שהוא יישאר "
              "תומך-החלטה ולא מחליף.")
add_kv_bullet(doc, "כוח ואחריותיות", "ריכוז נתונים ומודלים מעביר כוח לבעלי הפלטפורמה, ומעלה שאלת אחריות "
              "על המלצה שגויה.")
add_kv_bullet(doc, "פרטיות בקנה מידה ופגיעה במוניטין", "שליחת קורות-חיים ל-LLM חיצוני וצבירת נתונים; וכן "
              "סיכון שתיוג אנומליה שגוי יפגע במעסיק לגיטימי.")
add_kv_bullet(doc, "עלות סביבתית וחישובית", "הסקת LLM בקנה מידה נושאת עלות משאבים שיש לקחת בחשבון.")

add_lead(doc, "15.3 הערכה כוללת")
add_par(doc, "הפוטנציאל החיובי משמעותי — כלי תומך-החלטה שמנגיש, מייעל ומגן על מחפש העבודה — בתנאי שההטיות "
             "מנוטרות והאדם נשאר בשליטה. בולט שבחירות התכן של הפרויקט (הסברוּת, כנות, מסגור מייעץ ולא פוסל, "
             "ועיגון דטרמיניסטי) הן בדיוק המנגנונים הממתנים את הסיכונים המרכזיים — ולכן ההשפעה נטו צפויה "
             "להיות חיובית כאשר הפתרון מיושם באחריות.")

add_spacer(doc, 10)
add_par(doc, "— סוף התיעוד —", color=GREY, size=10)

# write into the project/ deliverables folder (relative to this script in benchmarking/)
_here = os.path.dirname(os.path.abspath(__file__))
out = os.path.normpath(os.path.join(_here, "..", "project", "Project_Data_Report.docx"))
doc.save(out)
print("Saved:", out)
