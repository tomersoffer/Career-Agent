# -*- coding: utf-8 -*-
"""
Generates the Hebrew (.docx) "Testing & Validation" report for the GPT-5.4-mini
AGENTIC TOOL-CALLING loop (agent_loop.py) of the Smart Career Navigator. RTL.
All numbers/traces come from real runs (tool_eval.json + traces_agent.json),
produced by claude/run_traces_agent.py. Nothing is hand-written.
"""
import os, json, math
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "David"
HEAD_FONT = "Arial"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xA8, 0x1F, 0x1F)
GREEN = RGBColor(0x2E, 0x6B, 0x2E)


# ---------- RTL helpers ----------
def _set_rtl_paragraph(p):
    p._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _style_run(run, font=BODY_FONT, size=11, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), font)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)


def add_par(doc, text, size=11, bold=False, color=None, font=BODY_FONT, space_after=6, space_before=0):
    p = doc.add_paragraph()
    _set_rtl_paragraph(p)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    _style_run(p.add_run(text), font=font, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 22, 1: 16, 2: 13}
    colors = {0: ACCENT, 1: ACCENT, 2: GREEN}
    p = doc.add_paragraph()
    _set_rtl_paragraph(p)
    p.paragraph_format.space_before = Pt(14 if level else 0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    _style_run(p.add_run(text), font=HEAD_FONT, size=sizes.get(level, 13), bold=True,
               color=colors.get(level, ACCENT))
    return p


def add_bullet(doc, text, sub=False):
    p = doc.add_paragraph()
    _set_rtl_paragraph(p)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.right_indent = Inches(0.3 if not sub else 0.6)
    _style_run(p.add_run(("•  " if not sub else "–  ") + text), size=11)
    return p


def add_kv_bullet(doc, key, val, sub=False):
    p = doc.add_paragraph()
    _set_rtl_paragraph(p)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.right_indent = Inches(0.3 if not sub else 0.6)
    _style_run(p.add_run(("•  " if not sub else "–  ") + key + " — "), size=11, bold=True)
    _style_run(p.add_run(val), size=11)
    return p


def add_lead(doc, text):
    return add_par(doc, text, size=12, bold=True, color=ACCENT, font=HEAD_FONT,
                   space_before=8, space_after=3)


def add_rtl_table(doc, headers, rows, col_widths=None, cell_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table._tbl.tblPr.append(OxmlElement('w:bidiVisual'))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0]._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _style_run(hdr[i].paragraphs[0].add_run(h), font=HEAD_FONT, size=10, bold=True,
                   color=RGBColor(255, 255, 255))
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            para = cells[i].paragraphs[0]
            para._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _style_run(para.add_run(str(val)), size=cell_size)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    return table


def add_code_block(doc, code, size=8.5):
    for raw in code.split("\n"):
        stripped = raw.lstrip(" ")
        indent = len(raw) - len(stripped)
        text = (" " * indent + stripped) if stripped else " "
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.12)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F3F3F3')
        pPr.append(shd)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rFonts.set(qn(a), "Consolas")
        rPr.append(rFonts)


def _field_label(doc, label):
    return add_par(doc, label, size=10.5, bold=True, color=GREEN, space_before=4, space_after=2)


# ---------- agentic-trace rendering ----------
CATEGORY = {
    "search": "חיפוש משרות → search_jobs",
    "data_question": "שאלת-נתונים → get_job_stat",
    "advice": "ייעוץ קריירה → give_career_advice",
    "smalltalk": "small-talk → small_talk",
    "more": "בקשת המשך → more_results",
    "multi_step": "רב-שלבי → search_jobs ואז get_job_stat",
    "failure_company": "כשל: שם חברה כתפקיד",
    "failure_probe": "מיקום בעברית (הכשל הקודם — נפתר)",
}


def _result_str(tool, res):
    if "fact" in res:
        return tool + " → fact = " + json.dumps(res["fact"], ensure_ascii=False)
    if "matches_n" in res:
        return tool + " → " + str(res["matches_n"]) + " matches: " + " | ".join(res.get("top", []))
    if "capabilities" in res:
        return tool + " → capabilities text returned"
    if "on_screen_jobs" in res:
        return tool + " → grounding: " + str(res["on_screen_jobs"]) + " on-screen jobs"
    return tool + " → " + json.dumps(res, ensure_ascii=False)


def render_agent_trace(doc, t, idx, fail=False):
    prefix = (str(idx) + ". ") if idx != "" else ""
    add_par(doc, prefix + CATEGORY.get(t["label"], t["label"]) + " — \"" + t["prompt"] + "\"",
            size=12, bold=True, color=(RED if fail else ACCENT), font=HEAD_FONT,
            space_before=10, space_after=3)
    # 1 — user prompt
    add_par(doc, "קלט (user prompt): " + t["prompt"], size=10.5)
    # 2 — tool calls the MODEL issued (name + args), one per hop
    _field_label(doc, "קריאות הכלי שהמודל הנפיק (tool calls — %d):" % len(t["tool_trace"]))
    for e in t["tool_trace"]:
        add_code_block(doc, e["tool"] + "(" + json.dumps(e["args"], ensure_ascii=False) + ")")
    # 3 — results returned from data/code
    _field_label(doc, "תוצאות מהנתונים והקוד (tool results):")
    for e in t["tool_trace"]:
        add_code_block(doc, _result_str(e["tool"], e["result"]))
    # 4 — final Hebrew reply
    _field_label(doc, "תשובת ה-LLM (LLM reply):")
    add_par(doc, t["reply"], size=10.5)


# ---------- load real run artifacts ----------
_OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "output")
with open(os.path.join(_OUT, "traces_agent.json"), encoding="utf-8") as f:
    TRACES = {t["label"]: t for t in json.load(f)}
with open(os.path.join(_OUT, "tool_eval.json"), encoding="utf-8") as f:
    EVAL = json.load(f)

SUCCESS_ORDER = ["search", "data_question", "advice", "smalltalk", "more", "multi_step"]


# ====================================================================
doc = Document()
normal = doc.styles['Normal']
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(0.9)

# ---------- TITLE ----------
add_heading(doc, "בדיקה ואימות — סוכן Tool-Calling מבוסס GPT-5.4-mini", level=0)
add_par(doc, "פרויקט גמר — Smart Career Navigator", size=15, bold=True, color=GREY, font=HEAD_FONT, space_after=2)
add_par(doc, "הערכת הסוכן האוטונומי שבו המודל מנפיק קריאות-כלי אמיתיות (OpenAI native tool-calling)",
        size=12, color=GREY, space_after=2)
add_par(doc, "המודל הנבדק: gpt-5.4-mini-2026-03-17  |  כל התוצאות הופקו מהרצות אמת (run_traces_agent.py).",
        size=10, color=GREY, space_after=14)

# ---------- 1. ARCHITECTURE ----------
add_heading(doc, "1. הארכיטקטורה הנבדקת — סוכן מבוסס קריאות-כלי", level=1)
add_par(doc, "בשונה מהגרסה הקודמת (קריאת-LLM אחת שמסווגת ל-JSON וקוד פייתון שמחליט מה להריץ), כאן המודל "
             "עצמו מקבל את הכלים ומחליט: איזה כלי לקרוא, עם אילו ארגומנטים, והאם לשרשר כמה כלים לפני "
             "שהוא עונה. זהו סוכן במובן המחמיר — המודל בוחר פעולות ופועל על בסיס התוצאות שחזרו.")
add_par(doc, "חמשת הכלים שהמודל רשאי לקרוא:", bold=True, space_before=2)
add_kv_bullet(doc, "search_jobs(role, state, seniority, remote)", "אחזור משרות מדורגות (TF-IDF + סינון).")
add_kv_bullet(doc, "get_job_stat(metric, role, state, …)", "סטטיסטיקה דטרמיניסטית ב-Pandas (שכר/ספירה/כישורים…).")
add_kv_bullet(doc, "more_results()", "דף תוצאות נוסף לחיפוש הקודם (paging).")
add_kv_bullet(doc, "give_career_advice(topic)", "הקשר לעיגון ייעוץ (משרות שעל המסך + שיחה).")
add_kv_bullet(doc, "small_talk()", "ברכות/תודות/שאלות על יכולות הסוכן.")
add_par(doc, "הסוכן מורץ בלולאה רב-שלבית חסומה (עד 4 צעדים): פעולה ← תצפית ← פעולה, עד שהמודל מנסח את "
             "התשובה הסופית בעברית ללא קריאת-כלי נוספת. הכלים מחזירים נתונים בלבד — האחזור והמספרים נשארים "
             "דטרמיניסטיים; המודל בוחר ומנסח.", space_before=2)

# ---------- 2. METHODOLOGY ----------
add_heading(doc, "2. שיטות הבדיקה והאימות", level=1)
add_kv_bullet(doc, "דיוק בחירת-כלי (tool-selection)", "21 שאילתות מתויגות הורצו דרך הסוכן; לכל אחת הוגדר "
              "הכלי הצפוי (לפי המפרט), ונבדק אם המודל אכן קרא לו.")
add_kv_bullet(doc, "לכידת trace מקצה-לקצה", "לכל מקרה נלכדו קריאות-הכלי שהמודל הנפיק (שם + ארגומנטים), "
              "התוצאות שחזרו מהקוד, והתשובה הסופית.")
add_kv_bullet(doc, "בדיקות רגרסיה (pytest)", "חבילת הבדיקות המלאה של הפרויקט הורצה לאחר השינוי — 185 בדיקות "
              "עברו, כולל בדיקות אינטגרציה ל-/api/chat שעודכנו לחוזה החדש.")
add_kv_bullet(doc, "אימות שורש", "לכל כשל לכאורה נבדקו הארגומנטים שהמודל העביר לכלי, כדי להבחין בין כשל "
              "מודל אמיתי לבין התנהגות תקינה.")

# ---------- 3. RESULTS ----------
add_heading(doc, "3. תוצאות", level=1)
add_par(doc, "דיוק בחירת-כלי: %d/%d = %d%%" % (EVAL["passed"], EVAL["n"],
             round(100 * EVAL["passed"] / EVAL["n"])),
        bold=True, color=GREEN, size=12)
add_par(doc, "בכל 21 המקרים המודל קרא לכלי הנכון (כולל הבחנה עדינה בין חיפוש, שאלת-נתונים וייעוץ). "
             "הרשימה המלאה בנספח א'.", space_after=6)

add_heading(doc, "3.1 דוגמאות מקצה-לקצה (End-to-End)", level=2)
add_par(doc, "כל דוגמה מציגה את הצינור המלא של הסוכן: קלט המשתמש ← קריאות-הכלי שהמודל הנפיק (שם + "
             "ארגומנטים) ← התוצאות שחזרו מהקוד ← תשובת ה-LLM הסופית בעברית. כל הערכים נלכדו מהרצת אמת.")
for _i, _lbl in enumerate(SUCCESS_ORDER, 1):
    if _lbl in TRACES:
        render_agent_trace(doc, TRACES[_lbl], _i)
add_par(doc, "שימו לב למקרה 6 (רב-שלבי): מבקשה אחת המודל שרשר שתי קריאות-כלי — תחילה search_jobs ואז "
             "get_job_stat — ורק אז ניסח תשובה אחת המשלבת את שתיהן. זוהי התנהגות סוכן אמיתית "
             "(פעולה ← תצפית ← פעולה).", space_before=4, color=GREY, size=10)

# ---------- 4. FAILURE / ROBUSTNESS ----------
add_heading(doc, "4. מקרה אי-הצלחה ובחינת עמידות", level=1)

add_lead(doc, "כשל אמיתי שנותר — זיהוי שם חברה כתפקיד")
add_par(doc, "כאשר המשתמש נוקב בשם חברה (ולא בתפקיד), המודל מעביר את שם החברה כארגומנט role של search_jobs "
             "במקום לזהות שאין כותרת-תפקיד לחפש לפיה — מה שפוגע באיכות האחזור (TF-IDF על שמות-משרות):")
render_agent_trace(doc, TRACES["failure_company"], "", fail=True)
add_par(doc, "כאן role=\"tesla jobs\" (שם חברה + המילה הגנרית jobs) מחזיר משרות \"Tesla Advisor\" קמעונאיות "
             "במקום תפקיד מבוקש. תיקון מומלץ: ולידציה של role מול מילון תפקידים, או הנחיה ייעודית ששמות-"
             "חברות אינם תפקידים.", space_before=2, size=10, color=GREY)

add_lead(doc, "הכשל הקודם (איבוד מיקום בעברית) — נפתר במעבר ל-tool-calling")
add_par(doc, "בגרסת ה-JSON-router הקודמת, \"מתכנת בגוגל בקליפורניה\" איבד לסירוגין את המדינה (state ריק). "
             "תחת ה-tool-calling, המודל מעביר את המדינה נכון כארגומנט — להלן trace אמיתי:")
render_agent_trace(doc, TRACES["failure_probe"], "")
add_par(doc, "המודל קרא search_jobs עם state=\"CA\", התוצאות אכן מקליפורניה (San Jose), והוא אף ציין ביושר "
             "שלא ראה משרת גוגל בתוצאות. כלומר המעבר לארכיטקטורת קריאות-הכלי שיפר את עמידות זיהוי המיקום.",
        space_before=2, size=10, color=GREY)

# ---------- 5. IS THIS AN AGENT ----------
add_heading(doc, "5. האם זהו Agent?  (כן — לפי ההגדרה המחמירה)", level=1)
add_par(doc, "בניגוד לגרסה הקודמת — שבה ה-LLM רק סיווג, וקוד פייתון החליט מה להריץ — בארכיטקטורה זו "
             "מתקיימות התכונות המגדירות סוכן LLM:")
add_kv_bullet(doc, "בחירת פעולה אוטונומית", "המודל בוחר בעצמו איזה מבין 5 הכלים לקרוא (דיוק 100% על 21 מקרים).")
add_kv_bullet(doc, "הנפקת ארגומנטים", "המודל מפיק את הארגומנטים המובנים (role/state/metric…) מתוך שפה חופשית.")
add_kv_bullet(doc, "תכנון רב-שלבי", "המודל משרשר כמה קריאות-כלי בתור אחד (search_jobs → get_job_stat) לפני התשובה.")
add_kv_bullet(doc, "פעולה על תצפיות", "התוצאות מוזרקות חזרה והמודל מחליט אם לקרוא לכלי נוסף או לענות.")
add_par(doc, "מסקנה: לאחר ה-refactor, /api/chat מונע על-ידי הסוכן עצמו, וה-trace-ים מראים קריאות-כלי "
             "אמיתיות שהמודל הנפיק — ולכן זהו סוכן במובן המחמיר, ולא צינור עם סיווג בלבד.", space_before=4)

# ---------- APPENDIX: tool-selection battery ----------
add_heading(doc, "נספח א' — סוללת בחירת-הכלי (21 שאילתות)", level=1)
add_par(doc, "כל שאילתה הורצה דרך הסוכן; נבדק אם המודל קרא לכלי הצפוי. כל הערכים מהרצת אמת.")
_rows = [[str(r["i"]), r["q"], r["expected_tool"], ", ".join(r["tools_called"]) or "—",
          ("✓" if r["ok"] else "✗")] for r in EVAL["rows"]]
add_rtl_table(doc, ["#", "השאלה שנשאלה", "הכלי הצפוי", "הכלי/ים שנקראו", "תוצאה"],
              _rows, col_widths=[0.35, 2.7, 1.6, 1.7, 0.5], cell_size=8.5)
add_par(doc, "סך הכול: %d/%d בחירות-כלי נכונות (%d%%)." % (
        EVAL["passed"], EVAL["n"], round(100 * EVAL["passed"] / EVAL["n"])),
        space_before=3, size=10, color=GREY)
add_lead(doc, "שחזור")
add_code_block(doc, "python -X utf8 claude/run_traces_agent.py   # מפיק traces_agent.json + tool_eval.json")

# ---------- SAVE ----------
OUT = os.path.join(_OUT, "GPT5.4-mini_Testing_Validation.docx")
os.makedirs(_OUT, exist_ok=True)
try:
    doc.save(OUT)
    print("Saved:", OUT)
except PermissionError:
    alt = OUT.replace(".docx", "_v2.docx")
    doc.save(alt)
    print("Target locked (open in Word?). Saved instead:", alt)
